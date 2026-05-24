import io
import logging
from pypdf import PdfReader
from pypdf._page import PageObject
from docx import Document

logger = logging.getLogger(__name__)

class Parser:
    @staticmethod
    def _extract_page_text(page: PageObject) -> str:
        """Extract text from a page with fallback methods for problematic PDFs."""
        # Try standard extraction first
        try:
            text = page.extract_text()
            if text:
                return text
        except KeyError as e:
            logger.warning(f"Standard text extraction failed: {e}")

        # Fallback: try extracting from content stream directly
        try:
            if "/Contents" in page:
                content = page.get_contents()
                if content:
                    # Try to get raw text operations from the content stream
                    text_parts = []
                    if hasattr(content, "get_data"):
                        data = content.get_data().decode("latin-1", errors="ignore")
                        # Extract text between parentheses (PDF text operators)
                        import re
                        matches = re.findall(r'\(([^)]*)\)', data)
                        text_parts.extend(matches)
                    if text_parts:
                        return " ".join(text_parts)
        except Exception as e:
            logger.warning(f"Fallback content stream extraction failed: {e}")

        # Last resort: try to extract any string annotations
        try:
            if "/Annots" in page:
                annots = page["/Annots"]
                text_parts = []
                for annot in annots:
                    annot_obj = annot.get_object()
                    if "/Contents" in annot_obj:
                        text_parts.append(str(annot_obj["/Contents"]))
                if text_parts:
                    return " ".join(text_parts)
        except Exception as e:
            logger.warning(f"Annotation extraction failed: {e}")

        return ""

    @staticmethod
    def from_pdf(pdf_bytes: bytes):
        # Wrap bytes in a file-like object
        bytes_stream = io.BytesIO(pdf_bytes)

        # Initialize the reader
        reader = PdfReader(bytes_stream)

        # Iterate and extract text from each page
        full_text = ""
        failed_pages = []
        for i, page in enumerate(reader.pages):
            page_text = Parser._extract_page_text(page)
            if page_text:
                full_text += page_text + "\n"
            else:
                failed_pages.append(i + 1)

        if failed_pages:
            logger.warning(f"Could not extract text from pages: {failed_pages}")

        return full_text

    @staticmethod
    def from_txt(txt_bytes: bytes) -> str:
        """Extract text from a TXT file with encoding detection."""
        # Try common encodings in order of likelihood
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                text = txt_bytes.decode(encoding)
                # Check for replacement characters that indicate wrong encoding
                if '\ufffd' not in text:
                    return text
            except (UnicodeDecodeError, LookupError):
                continue

        # Last resort: decode with errors ignored
        logger.warning("Could not detect encoding, using utf-8 with errors ignored")
        return txt_bytes.decode('utf-8', errors='ignore')
    
    @staticmethod
    def from_docx(docx_bytes: bytes) -> str:
        """Extract text from a DOCX file including paragraphs and tables."""
        bytes_stream = io.BytesIO(docx_bytes)
        doc = Document(bytes_stream)

        text_parts = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)